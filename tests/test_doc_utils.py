"""
test_doc_utils.py

测试doc_utils.py里的parse_document和split_text这两个核心函数。
这两个函数都是纯逻辑（不依赖数据库、不依赖网络、不依赖AI模型），
所以这里完全不需要用mock，直接测真实的函数行为就行——这是最"便宜"、最应该优先写的一类测试。
"""

import pytest
from doc_utils import parse_document, split_text


class TestSplitText:
    """测试文本切块逻辑，重点覆盖边界情况"""

    def test_empty_text_returns_empty_list(self):
        """空文本应该切出0块，而不是报错或者返回一个空字符串的块"""
        assert split_text("") == []

    def test_text_shorter_than_chunk_size_returns_single_chunk(self):
        """文本长度小于chunk_size时，应该整个作为一块，不需要切分"""
        text = "这是一段很短的文本"
        chunks = split_text(text, chunk_size=300, overlap=50)
        assert chunks == [text]

    def test_text_exactly_chunk_size_produces_redundant_tail_chunk(self):
        """
        文本长度正好等于chunk_size的边界情况。

        这里记录一个从写测试中发现的真实行为（不是想当然的假设）：
        循环用 `start < len(text)` 判断要不要继续切，但start每次只前进
        (chunk_size - overlap)这么多，不是chunk_size本身。
        所以即使第一块已经完整覆盖了全部文本，start前进后仍然小于文本长度，
        循环还会继续，多切出一个跟第一块结尾重叠、内容完全冗余的"尾巴块"。

        这不算丢数据的bug（顶多是多存了一份重复内容），但如果不希望
        出现这种冗余块，需要在split_text里加一个"已完整覆盖就提前退出"的判断。
        这里先如实记录当前行为，作为后续可选的优化点。
        """
        text = "a" * 300
        chunks = split_text(text, chunk_size=300, overlap=50)
        assert len(chunks) == 2
        assert chunks[0] == text
        assert chunks[1] == text[250:300]  # 冗余的尾巴块，内容完全被第一块包含

    def test_adjacent_chunks_have_overlap(self):
        """核心逻辑验证：相邻两块之间应该有overlap指定长度的重叠部分，
        这是为了避免在句子中间切断导致语义丢失"""
        text = "a" * 350
        chunks = split_text(text, chunk_size=300, overlap=50)
        assert len(chunks) == 2
        # 第二块应该从 250 (=300-50) 开始，而不是从300开始，验证确实有重叠
        assert chunks[1] == text[250:550]

    def test_custom_chunk_size_and_overlap_values(self):
        """用更小的数字手动验证切块的具体位置是否符合预期，逻辑上更容易肉眼核对"""
        text = "0123456789"
        chunks = split_text(text, chunk_size=4, overlap=1)
        # start=0: "0123" | start=3: "3456" | start=6: "6789" | start=9: "9"
        assert chunks == ["0123", "3456", "6789", "9"]


class TestParseDocument:
    """测试文档解析：格式路由 + 真实文件读取（用pytest的tmp_path生成临时测试文件）"""

    def test_txt_file_reads_correctly(self, tmp_path):
        file_path = tmp_path / "sample.txt"
        file_path.write_text("这是测试内容", encoding="utf-8")

        result = parse_document(str(file_path))

        assert result == "这是测试内容"

    def test_docx_file_reads_correctly(self, tmp_path):
        """用python-docx库真实创建一个docx文件，再验证我们自己的parse_document能读出来
        这是一个真实的"往返测试"（round-trip test），不需要mock"""
        from docx import Document

        doc = Document()
        doc.add_paragraph("第一段内容")
        doc.add_paragraph("第二段内容")
        file_path = tmp_path / "sample.docx"
        doc.save(str(file_path))

        result = parse_document(str(file_path))

        assert "第一段内容" in result
        assert "第二段内容" in result

    def test_pdf_file_with_no_extractable_text_returns_empty_string(self, tmp_path):
        """边界情况：一个没有文字内容的空白PDF页面，extract_text()会返回None，
        doc_utils.py里用了 `or ""` 兜底，验证这个兜底逻辑真的生效，不会报错"""
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        file_path = tmp_path / "sample.pdf"
        with open(file_path, "wb") as f:
            writer.write(f)

        result = parse_document(str(file_path))

        assert result == ""

    def test_unsupported_format_raises_value_error(self, tmp_path):
        """不支持的格式应该明确报错，而不是静默失败或者返回奇怪的结果"""
        file_path = tmp_path / "sample.xyz"
        file_path.write_text("内容", encoding="utf-8")

        with pytest.raises(ValueError, match="不支持的文件格式"):
            parse_document(str(file_path))
